from io import BytesIO

from docx import Document as WordDocument
from fastapi.testclient import TestClient


def create_course(client: TestClient) -> dict:
    response = client.post(
        "/courses",
        json={"name": "Document Processing", "code": "DOC101"},
    )
    assert response.status_code == 201
    return response.json()


def upload_document(
    client: TestClient, course_id: str, filename: str, content: bytes, mime_type: str
) -> dict:
    response = client.post(
        f"/courses/{course_id}/documents",
        data={"document_type": "READING"},
        files={"file": (filename, content, mime_type)},
    )
    assert response.status_code == 201
    return response.json()


def make_docx() -> bytes:
    document = WordDocument()
    document.add_heading("Vector Spaces", level=1)
    document.add_paragraph("A vector space is closed under addition and scalar multiplication.")
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def make_pdf(text: str) -> bytes:
    escaped_text = text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")
    content = f"BT /F1 12 Tf 72 720 Td ({escaped_text}) Tj ET".encode()
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        (
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>"
        ),
        b"<< /Length " + str(len(content)).encode() + b" >>\nstream\n" + content + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    pdf = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for number, obj in enumerate(objects, start=1):
        offsets.append(len(pdf))
        pdf.extend(f"{number} 0 obj\n".encode())
        pdf.extend(obj)
        pdf.extend(b"\nendobj\n")
    xref_offset = len(pdf)
    pdf.extend(f"xref\n0 {len(objects) + 1}\n".encode())
    pdf.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        pdf.extend(f"{offset:010d} 00000 n \n".encode())
    trailer = (
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
        f"startxref\n{xref_offset}\n%%EOF\n"
    )
    pdf.extend(trailer.encode())
    return bytes(pdf)


def process_and_get_chunks(client: TestClient, course_id: str, document_id: str) -> list[dict]:
    response = client.post(f"/courses/{course_id}/documents/{document_id}/process")
    assert response.status_code == 200
    assert response.json()["processingStatus"] == "READY"
    chunks_response = client.get(f"/courses/{course_id}/documents/{document_id}/chunks")
    assert chunks_response.status_code == 200
    return chunks_response.json()


def test_processes_plain_text_with_line_source_location(client: TestClient) -> None:
    course = create_course(client)
    document = upload_document(
        client, course["id"], "notes.txt", b"First line\nSecond line", "text/plain"
    )
    chunks = process_and_get_chunks(client, course["id"], document["id"])
    assert [chunk["sequence"] for chunk in chunks] == [0]
    assert chunks[0]["documentId"] == document["id"]
    assert chunks[0]["courseId"] == course["id"]
    assert chunks[0]["content"] == "First line\nSecond line"
    assert chunks[0]["pageNumber"] is None
    assert chunks[0]["sourceLocation"] == {"type": "lines", "start": 1, "end": 2}


def test_processes_markdown_with_heading_source(client: TestClient) -> None:
    course = create_course(client)
    document = upload_document(
        client,
        course["id"],
        "vectors.md",
        b"# Foundations\nVector addition is associative.\n\n## Bases\nA basis spans the space.",
        "text/markdown",
    )
    chunks = process_and_get_chunks(client, course["id"], document["id"])
    assert [chunk["section"] for chunk in chunks] == ["Foundations", "Bases"]
    assert chunks[0]["sourceLocation"] == {"type": "lines", "start": 2, "end": 2}
    assert chunks[1]["sourceLocation"] == {"type": "lines", "start": 5, "end": 5}


def test_processes_pdf_with_page_source_location(client: TestClient) -> None:
    course = create_course(client)
    document = upload_document(
        client, course["id"], "lecture.pdf", make_pdf("Gradient theorem"), "application/pdf"
    )
    chunks = process_and_get_chunks(client, course["id"], document["id"])
    assert chunks[0]["content"] == "Gradient theorem"
    assert chunks[0]["pageNumber"] == 1
    assert chunks[0]["sourceLocation"] == {
        "type": "page_lines",
        "start": 1,
        "end": 1,
        "page": 1,
    }


def test_processes_word_document_with_heading_source(client: TestClient) -> None:
    course = create_course(client)
    mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    document = upload_document(client, course["id"], "reading.docx", make_docx(), mime_type)
    chunks = process_and_get_chunks(client, course["id"], document["id"])
    assert chunks[0]["section"] == "Vector Spaces"
    assert "closed under addition" in chunks[0]["content"]
    assert chunks[0]["sourceLocation"] == {
        "type": "paragraphs",
        "start": 2,
        "end": 2,
    }


def test_repeat_processing_replaces_chunks_without_duplicates(client: TestClient) -> None:
    course = create_course(client)
    document = upload_document(
        client, course["id"], "repeat.txt", b"Repeatable content", "text/plain"
    )
    first_chunks = process_and_get_chunks(client, course["id"], document["id"])
    second_chunks = process_and_get_chunks(client, course["id"], document["id"])
    assert len(first_chunks) == len(second_chunks) == 1
    assert first_chunks[0]["id"] != second_chunks[0]["id"]
    assert second_chunks[0]["sequence"] == 0


def test_processing_failure_sets_failed_and_persists_no_chunks(client: TestClient) -> None:
    course = create_course(client)
    document = upload_document(
        client, course["id"], "broken.pdf", b"not a pdf", "application/pdf"
    )
    response = client.post(f"/courses/{course['id']}/documents/{document['id']}/process")
    assert response.status_code == 422
    documents = client.get(f"/courses/{course['id']}/documents").json()
    assert documents[0]["processingStatus"] == "FAILED"
    chunks = client.get(
        f"/courses/{course['id']}/documents/{document['id']}/chunks"
    ).json()
    assert chunks == []
